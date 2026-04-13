"""
=============================================================================
GENERADOR DE REPORTES BIOMÉTRICOS — FORMATO DOCX (Microsoft Word)
=============================================================================
Equivalente DOCX de las funciones generar_pdf / generar_pdf_persona
del módulo script.py.

Requiere:  pip install python-docx
=============================================================================
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ══════════════════════════════════════════════════════════════════════════
# PALETA DE COLORES (coincide con script.py)
# ══════════════════════════════════════════════════════════════════════════

_C_HEADER    = RGBColor(0x1a, 0x3a, 0x5c)   # Azul oscuro corporativo
_C_SUBHEADER = RGBColor(0x2e, 0x6d, 0xa4)   # Azul medio
_C_OK        = RGBColor(0xd4, 0xed, 0xda)   # Verde suave
_C_WARN      = RGBColor(0xff, 0xf3, 0xcd)   # Amarillo suave
_C_ERROR     = RGBColor(0xf8, 0xd7, 0xda)   # Rojo suave
_C_ALT       = RGBColor(0xf0, 0xf4, 0xf8)   # Gris muy suave
_C_TEXT      = RGBColor(0x21, 0x25, 0x29)   # Texto oscuro
_C_MUTED     = RGBColor(0x6c, 0x75, 0x7d)   # Texto gris
_C_WHITE     = RGBColor(0xff, 0xff, 0xff)
_C_BORDER    = RGBColor(0xde, 0xe2, 0xe6)


# ══════════════════════════════════════════════════════════════════════════
# UTILIDADES DE FORMATO
# ══════════════════════════════════════════════════════════════════════════

def _rgb_hex(rgb: RGBColor) -> str:
    """Convierte RGBColor a string hex sin '#' para uso en XML OOXML."""
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _set_cell_bg(cell, rgb: RGBColor):
    """Aplica color de fondo a una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), _rgb_hex(rgb))
    tcPr.append(shd)


def _set_cell_borders(cell, color: RGBColor = _C_BORDER):
    """Aplica bordes finos a una celda."""
    hex_c = _rgb_hex(color)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), hex_c)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _cell_text(cell, text: str, bold=False, color: RGBColor = None,
               size_pt: int = 9, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Escribe texto en una celda con formato."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int = 1,
                 color: RGBColor = _C_HEADER, size_pt: int = 13):
    """Agrega un párrafo de título con color."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    return p


def _add_rule(doc: Document, color: RGBColor = _C_SUBHEADER):
    """Agrega una línea horizontal usando borde inferior."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), _rgb_hex(color))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _setup_document(doc: Document):
    """Configura márgenes de página estándar A4."""
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)


def _tabla_portada(doc: Document, filas: list[list[str]]):
    """Tabla de metadatos para la portada (etiqueta + valor)."""
    tbl = doc.add_table(rows=len(filas), cols=2)
    tbl.style = 'Table Grid'
    col_widths = [Cm(5.5), Cm(11)]
    for i, (label, value) in enumerate(filas):
        row = tbl.rows[i]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
        bg = _C_ALT if i % 2 else _C_WHITE
        _set_cell_bg(row.cells[0], bg)
        _set_cell_bg(row.cells[1], bg)
        _set_cell_borders(row.cells[0])
        _set_cell_borders(row.cells[1])
        _cell_text(row.cells[0], label, bold=True, color=_C_SUBHEADER, size_pt=10)
        _cell_text(row.cells[1], value, color=_C_TEXT, size_pt=10)


def _tabla_datos(doc: Document, encabezados: list[str], filas: list[list],
                 col_widths: list = None, color_header: RGBColor = _C_HEADER,
                 color_novedad: RGBColor = None):
    """
    Crea una tabla con encabezado oscuro y filas alternas.
    color_novedad: si se especifica, todas las filas de datos se pintan con ese color.
    """
    n_cols = len(encabezados)
    tbl = doc.add_table(rows=1 + len(filas), cols=n_cols)
    tbl.style = 'Table Grid'

    # Encabezado
    hdr = tbl.rows[0]
    for j, txt in enumerate(encabezados):
        c = hdr.cells[j]
        _set_cell_bg(c, color_header)
        _set_cell_borders(c, _C_WHITE)
        _cell_text(c, txt, bold=True, color=_C_WHITE, size_pt=8,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        if col_widths:
            c.width = col_widths[j]

    # Filas de datos
    for i, fila in enumerate(filas):
        row = tbl.rows[i + 1]
        bg = color_novedad if color_novedad else (_C_ALT if i % 2 else _C_WHITE)
        for j, val in enumerate(fila):
            c = row.cells[j]
            _set_cell_bg(c, bg)
            _set_cell_borders(c, _C_BORDER)
            _cell_text(c, str(val) if val is not None else '', size_pt=8)
            if col_widths:
                c.width = col_widths[j]

    return tbl


# ══════════════════════════════════════════════════════════════════════════
# SECCIONES REUTILIZABLES
# ══════════════════════════════════════════════════════════════════════════

def _seccion_lista_tardanza(doc: Document, titulo: str, subtitle: str,
                             lista: list, color_bg: RGBColor):
    """Tabla de tardanza (severa o leve) con columnas: Persona | Llegada | Retraso | Programado."""
    if not lista:
        return
    _add_heading(doc, titulo, size_pt=10, color=_C_SUBHEADER)
    p = doc.add_paragraph(subtitle)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = _C_MUTED
    p.paragraph_format.space_after = Pt(3)

    encabezados = ["Persona", "Llegada", "Retraso (min)", "Hora programada"]
    filas = [
        [r["nombre"], r.get("hora", "—"),
         str(r.get("retraso", "")), r.get("programado", "—")]
        for r in lista
    ]
    widths = [Cm(7), Cm(2.5), Cm(3), Cm(3.5)]
    _tabla_datos(doc, encabezados, filas, col_widths=widths,
                 color_header=_C_SUBHEADER, color_novedad=color_bg)
    doc.add_paragraph()


def _seccion_almuerzo(doc: Document, lista: list):
    """Tabla de excesos de almuerzo."""
    if not lista:
        return
    _add_heading(doc, "Exceso de Almuerzo", size_pt=10, color=_C_SUBHEADER)
    encabezados = ["Persona", "Duración (min)", "Exceso (min)", "Salida alm.", "Regreso alm."]
    filas = [
        [r["nombre"],
         str(r.get("duracion", "—")),
         str(r.get("exceso", "—")),
         r.get("salida_alm", "—"),
         r.get("regreso_alm", "—")]
        for r in lista
    ]
    widths = [Cm(6), Cm(3), Cm(3), Cm(2.5), Cm(3)]
    _tabla_datos(doc, encabezados, filas, col_widths=widths,
                 color_header=_C_SUBHEADER, color_novedad=_C_WARN)
    doc.add_paragraph()


def _seccion_incompletos(doc: Document, lista: list):
    """Tabla de registros incompletos/anómalos."""
    if not lista:
        return
    _add_heading(doc, "Registros Incompletos / Anómalos", size_pt=10, color=_C_SUBHEADER)
    encabezados = ["Persona", "Detalle registros", "Observaciones"]
    filas = [
        [r["nombre"],
         r.get("detalle_registros", "—"),
         "; ".join(r.get("observaciones", []))]
        for r in lista
    ]
    widths = [Cm(5), Cm(5.5), Cm(6)]
    _tabla_datos(doc, encabezados, filas, col_widths=widths,
                 color_header=_C_SUBHEADER, color_novedad=_C_ERROR)
    doc.add_paragraph()


def _seccion_salidas_anticipadas(doc: Document, titulo: str, lista: list, color_bg: RGBColor):
    if not lista:
        return
    _add_heading(doc, titulo, size_pt=10, color=_C_SUBHEADER)
    encabezados = ["Persona", "Hora salida", "Anticipación (min)", "Hora programada"]
    filas = [
        [r["nombre"], r.get("hora", "—"),
         str(r.get("retraso", "")), r.get("programado", "—")]
        for r in lista
    ]
    widths = [Cm(7), Cm(2.5), Cm(3.5), Cm(3.5)]
    _tabla_datos(doc, encabezados, filas, col_widths=widths,
                 color_header=_C_SUBHEADER, color_novedad=color_bg)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════
# MODO GENERAL: generar_docx
# ══════════════════════════════════════════════════════════════════════════

def generar_docx(
    ruta_salida: str,
    analisis_por_dia: dict,
    log_duplicados: list[dict],
    config: dict,
    nombre_archivo_origen: str,
    filtros: dict = None,
    sin_horario: list = None,
):
    """
    Genera el reporte DOCX equivalente a generar_pdf().
    Mismos parámetros y estructura de datos.
    """
    if filtros is None:
        filtros = {}
    if sin_horario is None:
        sin_horario = []

    _F = {
        "mostrar_tardanza_leve":     filtros.get("mostrar_tardanza_leve",     True),
        "mostrar_tardanza_severa":   filtros.get("mostrar_tardanza_severa",   True),
        "mostrar_almuerzo":          filtros.get("mostrar_almuerzo",          True),
        "mostrar_incompletos":       filtros.get("mostrar_incompletos",       True),
        "mostrar_salida_anticipada": filtros.get("mostrar_salida_anticipada", True),
        "mostrar_todos_los_dias":    filtros.get("mostrar_todos_los_dias",    False),
    }

    nombre_sistema     = os.getenv("NOMBRE_SISTEMA",     "Informes Biométricos")
    nombre_institucion = os.getenv("NOMBRE_INSTITUCION", "ISTPET")

    doc = Document()
    _setup_document(doc)

    # ── Portada ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_titulo.add_run(nombre_sistema.upper())
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = _C_HEADER

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run(f"Control Biométrico · {nombre_institucion.upper()}")
    r2.font.size = Pt(12)
    r2.font.color.rgb = _C_SUBHEADER

    doc.add_paragraph()
    _add_rule(doc, _C_HEADER)

    # Resumen de fechas
    fechas = sorted(analisis_por_dia.keys())
    fecha_str = "—"
    if fechas:
        if hasattr(fechas[0], "strftime"):
            fecha_str = f"{fechas[0].strftime('%d/%m/%Y')} — {fechas[-1].strftime('%d/%m/%Y')}"
        else:
            fecha_str = f"{fechas[0]} — {fechas[-1]}"

    total_dias = len(analisis_por_dia)
    total_personas = max(
        (d["resumen"].get("total_personas", 0) for d in analisis_por_dia.values()),
        default=0
    )

    meta = [
        ["Archivo origen:", nombre_archivo_origen],
        ["Generado el:", datetime.now().strftime("%d/%m/%Y %H:%M")],
        ["Período:", fecha_str],
        ["Días con registros:", str(total_dias)],
        ["Personas analizadas:", str(total_personas)],
    ]
    _tabla_portada(doc, meta)
    doc.add_page_break()

    # ── Resumen ejecutivo ─────────────────────────────────────────────────
    _add_heading(doc, "RESUMEN EJECUTIVO DEL PERÍODO", color=_C_HEADER, size_pt=13)
    _add_rule(doc)

    total_ts = sum(d["resumen"].get("tardanza_severa", 0) for d in analisis_por_dia.values())
    total_tl = sum(d["resumen"].get("tardanza_leve", 0)   for d in analisis_por_dia.values())
    total_alm = sum(d["resumen"].get("almuerzo_largo", 0) for d in analisis_por_dia.values())
    total_inc = sum(d["resumen"].get("incompletos", 0)    for d in analisis_por_dia.values())
    total_sa  = sum(
        d["resumen"].get("salida_anticipada_severa", 0) + d["resumen"].get("salida_anticipada_leve", 0)
        for d in analisis_por_dia.values()
    )

    enc_res = ["Indicador", "Total"]
    filas_res = [
        ["Tardanzas severas",    str(total_ts)],
        ["Tardanzas leves",      str(total_tl)],
        ["Excesos de almuerzo",  str(total_alm)],
        ["Registros incompletos",str(total_inc)],
        ["Salidas anticipadas",  str(total_sa)],
    ]
    _tabla_datos(doc, enc_res, filas_res, col_widths=[Cm(10), Cm(3)])
    doc.add_paragraph()

    # ── Selección de días a mostrar ───────────────────────────────────────
    _claves_activas = []
    if _F["mostrar_tardanza_leve"]:    _claves_activas.append("tardanza_leve")
    if _F["mostrar_tardanza_severa"]:  _claves_activas.append("tardanza_severa")
    if _F["mostrar_almuerzo"]:         _claves_activas.append("almuerzo_largo")
    if _F["mostrar_incompletos"]:      _claves_activas.append("incompletos")
    if _F["mostrar_salida_anticipada"]:
        _claves_activas.append("salida_anticipada_leve")
        _claves_activas.append("salida_anticipada_severa")

    if _F["mostrar_todos_los_dias"]:
        dias_a_mostrar = sorted(analisis_por_dia.keys())
    else:
        dias_a_mostrar = [
            d for d in sorted(analisis_por_dia.keys())
            if any(analisis_por_dia[d]["resumen"].get(k, 0) > 0 for k in _claves_activas)
        ]

    if not dias_a_mostrar:
        p = doc.add_paragraph(
            f"✓ Sin novedades en el período consultado ({total_dias} días con registros)."
        )
        p.runs[0].font.color.rgb = RGBColor(0x15, 0x52, 0x24)
    else:
        for dia in dias_a_mostrar:
            datos_dia = analisis_por_dia[dia]
            res = datos_dia["resumen"]

            if hasattr(dia, "strftime"):
                dia_str = dia.strftime("%A %d/%m/%Y").capitalize()
            else:
                dia_str = str(dia)

            _add_heading(doc, f"DÍA: {dia_str.upper()}", color=_C_HEADER, size_pt=11)
            _add_rule(doc)

            # Resumen del día
            partes = []
            if res.get("tardanza_severa"): partes.append(f"{res['tardanza_severa']} tard. severas")
            if res.get("tardanza_leve"):   partes.append(f"{res['tardanza_leve']} tard. leves")
            if res.get("almuerzo_largo"):  partes.append(f"{res['almuerzo_largo']} exc. almuerzo")
            if res.get("incompletos"):     partes.append(f"{res['incompletos']} anómalos")
            sa_t = res.get("salida_anticipada_severa", 0) + res.get("salida_anticipada_leve", 0)
            if sa_t: partes.append(f"{sa_t} sal. anticipadas")
            resumen_txt = "  ·  ".join(partes) if partes else "Sin novedades"
            p_res = doc.add_paragraph(resumen_txt)
            p_res.runs[0].font.size = Pt(9)
            p_res.runs[0].font.color.rgb = _C_MUTED
            p_res.paragraph_format.space_after = Pt(6)

            if _F["mostrar_tardanza_severa"]:
                _seccion_lista_tardanza(doc, "Tardanzas Severas",
                                        "Llegadas con más de 5 minutos de retraso",
                                        datos_dia.get("tardanza_severa", []), _C_ERROR)
            if _F["mostrar_tardanza_leve"]:
                _seccion_lista_tardanza(doc, "Tardanzas Leves",
                                        "Llegadas con 1 a 5 minutos de retraso",
                                        datos_dia.get("tardanza_leve", []), _C_WARN)
            if _F["mostrar_almuerzo"]:
                _seccion_almuerzo(doc, datos_dia.get("almuerzo_largo", []))
            if _F["mostrar_incompletos"]:
                _seccion_incompletos(doc, datos_dia.get("registros_incompletos", []))
            if _F["mostrar_salida_anticipada"]:
                _seccion_salidas_anticipadas(doc, "Salidas Anticipadas Severas",
                                             datos_dia.get("salida_anticipada_severa", []), _C_ERROR)
                _seccion_salidas_anticipadas(doc, "Salidas Anticipadas Leves",
                                             datos_dia.get("salida_anticipada_leve", []), _C_WARN)

    # ── Personas sin horario ──────────────────────────────────────────────
    if sin_horario:
        doc.add_page_break()
        _add_heading(doc, "PERSONAS SIN HORARIO ASIGNADO", color=_C_HEADER)
        _add_rule(doc)
        filas_sh = [[str(i + 1), n] for i, n in enumerate(sin_horario)]
        _tabla_datos(doc, ["#", "Nombre"], filas_sh, col_widths=[Cm(1.5), Cm(15)])
        doc.add_paragraph()

    # ── Log de duplicados ─────────────────────────────────────────────────
    if log_duplicados:
        doc.add_page_break()
        _add_heading(doc, "LOG DE REGISTROS DUPLICADOS", color=_C_HEADER)
        _add_rule(doc)
        enc_dup = ["Persona", "Fecha", "Hora 1", "Hora 2", "Diferencia (min)"]
        filas_dup = [
            [d.get("nombre", "—"), str(d.get("fecha", "—")),
             str(d.get("hora1", "—")), str(d.get("hora2", "—")),
             str(d.get("diferencia_min", "—"))]
            for d in log_duplicados
        ]
        _tabla_datos(doc, enc_dup, filas_dup,
                     col_widths=[Cm(5), Cm(3), Cm(2.5), Cm(2.5), Cm(3.5)])

    # ── Pie de página ─────────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_f = fp.add_run(
            f"Generado por {nombre_sistema} · {nombre_institucion} · "
            f"{datetime.now().strftime('%d/%m/%Y %H:%M')} — "
            "Documento de uso interno. No constituye acto administrativo."
        )
        run_f.font.size = Pt(7)
        run_f.font.color.rgb = _C_MUTED

    doc.save(ruta_salida)
    print(f"\n✅ Reporte DOCX generado: {ruta_salida}\n")


# ══════════════════════════════════════════════════════════════════════════
# MODO POR PERSONA: generar_docx_persona
# ══════════════════════════════════════════════════════════════════════════

def generar_docx_persona(
    ruta_salida: str,
    analisis_persona: dict,
    config: dict,
    nombre_archivo_origen: str,
    filtros: dict = None,
    sin_horario: list = None,
):
    """
    Genera el reporte DOCX equivalente a generar_pdf_persona().
    Mismos parámetros y estructura de datos.
    """
    if filtros is None:
        filtros = {}
    if sin_horario is None:
        sin_horario = []

    _F = {
        "mostrar_ausencias":         filtros.get("mostrar_ausencias",         True),
        "mostrar_tardanza_severa":   filtros.get("mostrar_tardanza_severa",   True),
        "mostrar_tardanza_leve":     filtros.get("mostrar_tardanza_leve",     True),
        "mostrar_almuerzo":          filtros.get("mostrar_almuerzo",          True),
        "mostrar_incompletos":       filtros.get("mostrar_incompletos",       True),
        "mostrar_salida_anticipada": filtros.get("mostrar_salida_anticipada", True),
        "mostrar_todos_los_dias":    filtros.get("mostrar_todos_los_dias",    False),
        "columna_tiempo_dentro":     filtros.get("columna_tiempo_dentro",     False),
        "verificar_horas":           filtros.get("verificar_horas",           False),
        "mostrar_tiempo_extra":      filtros.get("mostrar_tiempo_extra",      False),
    }

    nombre_sistema     = os.getenv("NOMBRE_SISTEMA",     "Informes Biométricos")
    nombre_institucion = os.getenv("NOMBRE_INSTITUCION", "ISTPET")

    doc = Document()
    _setup_document(doc)

    # ── Portada ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_titulo.add_run(f"{nombre_sistema.upper()} (POR PERSONA)")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = _C_HEADER

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run(f"Control Biométrico · {nombre_institucion.upper()}")
    r2.font.size = Pt(12)
    r2.font.color.rgb = _C_SUBHEADER

    doc.add_paragraph()
    _add_rule(doc, _C_HEADER)

    if len(analisis_persona) == 1:
        persona_nombre = list(analisis_persona.keys())[0]
        meta = [
            ["Archivo origen:",    nombre_archivo_origen],
            ["Generado el:",       datetime.now().strftime("%d/%m/%Y %H:%M")],
            ["Persona analizada:", persona_nombre],
        ]
    else:
        meta = [
            ["Archivo origen:",      nombre_archivo_origen],
            ["Generado el:",         datetime.now().strftime("%d/%m/%Y %H:%M")],
            ["Personas analizadas:", str(len(analisis_persona))],
        ]
    _tabla_portada(doc, meta)
    doc.add_page_break()

    # ── Resumen general ───────────────────────────────────────────────────
    _add_heading(doc, "RESUMEN GENERAL", color=_C_HEADER, size_pt=13)
    _add_rule(doc)

    enc_res = ["Persona", "Días", "Ausencias", "Tard. Sev.",
               "Tard. Lev.", "Sal. Ant.", "Exc. Alm.", "Anóm.", "Justif."]
    filas_resumen = []
    for nombre in sorted(analisis_persona.keys()):
        r_res = analisis_persona[nombre]["resumen"]
        tot_sa = r_res.get("salida_anticipada_severa", 0) + r_res.get("salida_anticipada_leve", 0)
        filas_resumen.append([
            nombre,
            str(r_res.get("total_dias",      0)),
            str(r_res.get("ausencias",       0)),
            str(r_res.get("tardanza_severa", 0)),
            str(r_res.get("tardanza_leve",   0)),
            str(tot_sa),
            str(r_res.get("almuerzo_largo",  0)),
            str(r_res.get("incompletos",     0)),
            str(r_res.get("justificadas",    0)),
        ])
    widths_res = [Cm(4.5), Cm(1.2), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.5), Cm(1.5)]
    _tabla_datos(doc, enc_res, filas_resumen, col_widths=widths_res)
    doc.add_paragraph()

    # ── Sección detallada por persona ─────────────────────────────────────
    for nombre in sorted(analisis_persona.keys()):
        datos_p = analisis_persona[nombre]
        r_res   = datos_p["resumen"]
        dias    = datos_p.get("dias", [])

        _add_heading(doc, f"  {nombre.upper()}", color=_C_HEADER, size_pt=11)
        _add_rule(doc)

        partes = []
        if r_res.get("ausencias",       0): partes.append(f"{r_res['ausencias']} ausencias")
        if r_res.get("tardanza_severa", 0): partes.append(f"{r_res['tardanza_severa']} tard. severas")
        if r_res.get("tardanza_leve",   0): partes.append(f"{r_res['tardanza_leve']} tard. leves")
        if r_res.get("almuerzo_largo",  0): partes.append(f"{r_res['almuerzo_largo']} exc. almuerzo")
        if r_res.get("incompletos",     0): partes.append(f"{r_res['incompletos']} anómalos")
        if r_res.get("justificadas",    0): partes.append(f"{r_res['justificadas']} justificadas")
        resumen_txt = "  ·  ".join(partes) if partes else "Sin novedades"
        p_res_p = doc.add_paragraph(resumen_txt)
        p_res_p.runs[0].font.size = Pt(9)
        p_res_p.runs[0].font.color.rgb = _C_MUTED

        if datos_p.get("sin_novedades") or not dias:
            total_d = r_res.get("total_dias", 0)
            sufijo  = "día analizado" if total_d == 1 else "días analizados"
            p_ok = doc.add_paragraph(
                f"✓ Sin novedades en el período consultado ({total_d} {sufijo} con registros)."
            )
            p_ok.runs[0].font.color.rgb = RGBColor(0x15, 0x52, 0x24)
            doc.add_paragraph()
            continue

        # Ausencias
        if _F["mostrar_ausencias"]:
            ausentes = [d for d in dias if d.get("estado") == "ausente"]
            if ausentes:
                _add_heading(doc, "Ausencias", size_pt=10, color=_C_SUBHEADER)
                enc_aus = ["Fecha", "Justificación", "Motivo"]
                filas_aus = []
                for d in ausentes:
                    j = d.get("justificacion") or {}
                    motivo = j.get("motivo", "—") if isinstance(j, dict) else "—"
                    justificado = "Sí" if j else "No"
                    filas_aus.append([
                        str(d["fecha"].strftime("%d/%m/%Y") if hasattr(d["fecha"], "strftime") else d["fecha"]),
                        justificado,
                        motivo,
                    ])
                _tabla_datos(doc, enc_aus, filas_aus,
                             col_widths=[Cm(3), Cm(2.5), Cm(10)],
                             color_header=_C_SUBHEADER, color_novedad=_C_ERROR)
                doc.add_paragraph()

        # Tardanzas severas
        if _F["mostrar_tardanza_severa"]:
            severas = [d for d in dias if d.get("estado") == "severa"]
            if severas:
                _add_heading(doc, "Tardanzas Severas", size_pt=10, color=_C_SUBHEADER)
                _seccion_detalle_persona(doc, severas, _F["columna_tiempo_dentro"], _C_ERROR)

        # Tardanzas leves
        if _F["mostrar_tardanza_leve"]:
            leves = [d for d in dias if d.get("estado") == "leve"
                     and any("Tardanza leve" in str(o) for o in d.get("observaciones", []))]
            if leves:
                _add_heading(doc, "Tardanzas Leves", size_pt=10, color=_C_SUBHEADER)
                _seccion_detalle_persona(doc, leves, _F["columna_tiempo_dentro"], _C_WARN)

        # Excesos de almuerzo
        if _F["mostrar_almuerzo"]:
            alm_dias = [d for d in dias
                        if d.get("almuerzo_exceso") is not None and d.get("almuerzo_exceso", 0) > 0]
            if alm_dias:
                _add_heading(doc, "Excesos de Almuerzo", size_pt=10, color=_C_SUBHEADER)
                enc_alm = ["Fecha", "Llegada", "Salida", "Exc. alm. (min)", "Sal. alm.", "Reg. alm."]
                filas_alm = [
                    [_fmt_fecha(d["fecha"]), d.get("llegada", "—"), d.get("salida", "—"),
                     str(d.get("almuerzo_exceso", "—")),
                     d.get("almuerzo_salida", "—"), d.get("almuerzo_regreso", "—")]
                    for d in alm_dias
                ]
                _tabla_datos(doc, enc_alm, filas_alm,
                             col_widths=[Cm(2.8), Cm(2), Cm(2), Cm(3), Cm(2.5), Cm(2.5)],
                             color_header=_C_SUBHEADER, color_novedad=_C_WARN)
                doc.add_paragraph()

        # Registros incompletos
        if _F["mostrar_incompletos"]:
            inc_dias = [d for d in dias if d.get("estado") == "incompleto"]
            if inc_dias:
                _add_heading(doc, "Registros Incompletos / Anómalos", size_pt=10, color=_C_SUBHEADER)
                _seccion_detalle_persona(doc, inc_dias, _F["columna_tiempo_dentro"], _C_ERROR)

        # Salidas anticipadas
        if _F["mostrar_salida_anticipada"]:
            sal_ant = [d for d in dias if any("Salida ant." in str(o) for o in d.get("observaciones", []))]
            if sal_ant:
                _add_heading(doc, "Salidas Anticipadas", size_pt=10, color=_C_SUBHEADER)
                _seccion_detalle_persona(doc, sal_ant, _F["columna_tiempo_dentro"], _C_WARN)

        # Verificación de horas de contrato
        if (_F["verificar_horas"] or _F["mostrar_tiempo_extra"]) and r_res.get("horas_contrato_tipo"):
            _seccion_horas_contrato_docx(doc, r_res)

        # Detalle cronológico completo
        if _F["mostrar_todos_los_dias"] and dias:
            _add_heading(doc, "Detalle Cronológico (Todos los días)", size_pt=10, color=_C_SUBHEADER)
            _seccion_detalle_persona(doc, dias, _F["columna_tiempo_dentro"], None)

        doc.add_paragraph()

    # ── Personas sin horario ──────────────────────────────────────────────
    if sin_horario:
        doc.add_page_break()
        _add_heading(doc, "PERSONAS SIN HORARIO ASIGNADO", color=_C_HEADER)
        _add_rule(doc)
        filas_sh = [[str(i + 1), n] for i, n in enumerate(sin_horario)]
        _tabla_datos(doc, ["#", "Nombre"], filas_sh, col_widths=[Cm(1.5), Cm(15)])

    # ── Pie de página ─────────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_f = fp.add_run(
            f"Generado por {nombre_sistema} · {nombre_institucion} · "
            f"{datetime.now().strftime('%d/%m/%Y %H:%M')} — "
            "Documento de uso interno. No constituye acto administrativo."
        )
        run_f.font.size = Pt(7)
        run_f.font.color.rgb = _C_MUTED

    doc.save(ruta_salida)
    print(f"\n✅ Reporte DOCX generado: {ruta_salida}\n")


# ══════════════════════════════════════════════════════════════════════════
# HELPERS INTERNOS PARA MODO PERSONA
# ══════════════════════════════════════════════════════════════════════════

def _fmt_fecha(fecha) -> str:
    if hasattr(fecha, "strftime"):
        return fecha.strftime("%d/%m/%Y")
    return str(fecha)


def _seccion_detalle_persona(doc: Document, dias: list, col_tiempo_dentro: bool,
                              color_novedad: RGBColor):
    """Tabla de detalle por día para un grupo de días de una persona."""
    if col_tiempo_dentro:
        encabezados = ["Fecha", "Estado", "Llegada", "Salida", "H. prog.", "T. dentro", "Observaciones"]
        widths = [Cm(2.5), Cm(2), Cm(1.8), Cm(1.8), Cm(1.8), Cm(2), Cm(5)]
    else:
        encabezados = ["Fecha", "Estado", "Llegada", "Salida", "H. prog.", "Observaciones"]
        widths = [Cm(2.5), Cm(2), Cm(1.8), Cm(1.8), Cm(1.8), Cm(7)]

    filas = []
    for d in dias:
        obs = "; ".join(str(o) for o in d.get("observaciones", []))
        fila = [
            _fmt_fecha(d["fecha"]),
            d.get("estado", "—"),
            d.get("llegada", "—") or "—",
            d.get("salida",  "—") or "—",
            d.get("hora_programada", "—") or "—",
        ]
        if col_tiempo_dentro:
            td = d.get("tiempo_dentro")
            fila.append(f"{td} min" if td is not None else "—")
        fila.append(obs or "—")
        filas.append(fila)

    _tabla_datos(doc, encabezados, filas, col_widths=widths,
                 color_header=_C_SUBHEADER, color_novedad=color_novedad)
    doc.add_paragraph()


def _seccion_horas_contrato_docx(doc: Document, resumen: dict):
    """Tabla de cumplimiento de horas de contrato."""
    _add_heading(doc, "Verificación de Horas de Contrato", size_pt=10, color=_C_SUBHEADER)
    tipo = resumen.get("horas_contrato_tipo", "—")
    horas_prog = resumen.get("horas_contrato_programadas", 0)
    horas_real = resumen.get("horas_reales_trabajadas", 0)
    diferencia = resumen.get("diferencia_horas", 0)
    color_dif = _C_OK if diferencia >= 0 else _C_ERROR

    enc = ["Tipo contrato", "Horas programadas", "Horas reales", "Diferencia"]
    filas = [[tipo, f"{horas_prog:.1f} h", f"{horas_real:.1f} h", f"{diferencia:+.1f} h"]]
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = 'Table Grid'
    widths = [Cm(4), Cm(4), Cm(4), Cm(4)]
    for j, txt in enumerate(enc):
        c = tbl.rows[0].cells[j]
        c.width = widths[j]
        _set_cell_bg(c, _C_SUBHEADER)
        _set_cell_borders(c, _C_WHITE)
        _cell_text(c, txt, bold=True, color=_C_WHITE, size_pt=8,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    for j, val in enumerate(filas[0]):
        c = tbl.rows[1].cells[j]
        c.width = widths[j]
        bg = color_dif if j == 3 else _C_ALT
        _set_cell_bg(c, bg)
        _set_cell_borders(c, _C_BORDER)
        _cell_text(c, val, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
